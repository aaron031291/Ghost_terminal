
"""Teach-and-Train Interface for Grace.

This module handles learning from various document types and integrating knowledge.
"""

import os
import logging
import re
import json
from typing import Dict, Any, Optional, List, Union, Tuple
from pathlib import Path
import docx
import markdown
import yaml

from grace.config import Config

logger = logging.getLogger("grace.training")


class DocumentParser:
    """Parser for various document types."""

    def __init__(self):
        """Initialize document parser."""
        self.parsers = {
            ".md": self.parse_markdown,
            ".markdown": self.parse_markdown,
            ".txt": self.parse_text,
            ".py": self.parse_python,
            ".js": self.parse_javascript,
            ".docx": self.parse_docx,
            ".json": self.parse_json,
            ".yaml": self.parse_yaml,
            ".yml": self.parse_yaml,
        }
        
    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """Parse file based on extension.
        
        Args:
            file_path: Path to file
            
        Returns:
            Dictionary with parsed content and metadata
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext not in self.parsers:
            logger.warning(f"No parser available for {ext} files")
            return self.parse_text(file_path)
            
        return self.parsers[ext](file_path)
        
    def parse_text(self, file_path: str) -> Dict[str, Any]:
        """Parse plain text file.
        
        Args:
            file_path: Path to text file
            
        Returns:
            Dictionary with parsed content
        """
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
                
            return {
                "type": "text",
                "content": content,
                "metadata": {
                    "file_path": file_path,
                    "file_size": os.path.getsize(file_path),
                },
            }
        except Exception as e:
            logger.error(f"Error parsing text file {file_path}: {str(e)}")
            return {
                "type": "text",
                "content": "",
                "error": str(e),
                "metadata": {
                    "file_path": file_path,
                },
            }
            
    def parse_markdown(self, file_path: str) -> Dict[str, Any]:
        """Parse Markdown file.
        
        Args:
            file_path: Path to Markdown file
            
        Returns:
            Dictionary with parsed content and structure
        """
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
                
            # Extract front matter if present
            front_matter = {}
            content_without_front_matter = content
            
            # Check for YAML front matter
            front_matter_match = re.match(r"^---\s*\n(.*?)\n---\s*\n", content, re.DOTALL)
            if front_matter_match:
                try:
                    front_matter = yaml.safe_load(front_matter_match.group(1))
                    content_without_front_matter = content[front_matter_match.end():]
                except Exception as e:
                    logger.warning(f"Error parsing front matter in {file_path}: {str(e)}")
                    
            # Extract headings and structure
            headings = []
            for match in re.finditer(r"^(#{1,6})\s+(.+)$", content, re.MULTILINE):
                level = len(match.group(1))
                text = match.group(2).strip()
                headings.append({
                    "level": level,
                    "text": text,
                    "position": match.start(),
                })
                
            # Extract code blocks
            code_blocks = []
            for match in re.finditer(r"```(\w*)\n(.*?)```", content, re.DOTALL):
                language = match.group(1) or "text"
                code = match.group(2)
                code_blocks.append({
                    "language": language,
                    "code": code,
                    "position": match.start(),
                })
                
            return {
                "type": "markdown",
                "content": content,
                "content_without_front_matter": content_without_front_matter,
                "front_matter": front_matter,
                "headings": headings,
                "code_blocks": code_blocks,
                "metadata": {
                    "file_path": file_path,
                    "file_size": os.path.getsize(file_path),
                },
            }
        except Exception as e:
            logger.error(f"Error parsing Markdown file {file_path}: {str(e)}")
            return {
                "type": "markdown",
                "content": "",
                "error": str(e),
                "metadata": {
                    "file_path": file_path,
                },
            }
            
    def parse_python(self, file_path: str) -> Dict[str, Any]:
        """Parse Python file.
        
        Args:
            file_path: Path to Python file
            
        Returns:
            Dictionary with parsed content and structure
        """
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
                
            # Extract docstring
            docstring = None
            docstring_match = re.match(r'"""(.*?)"""', content, re.DOTALL)
            if docstring_match:
                docstring = docstring_match.group(1).strip()
                
            # Extract imports
            imports = []
            for match in re.finditer(r"^(?:from\s+(\S+)\s+)?import\s+(.+)$", content, re.MULTILINE):
                from_module = match.group(1)
                import_items = match.group(2).split(",")
                
                for item in import_items:
                    item = item.strip()
                    if " as " in item:
                        name, alias = item.split(" as ")
                        imports.append({
                            "from": from_module,
                            "import": name.strip(),
                            "as": alias.strip(),
                        })
                    else:
                        imports.append({
                            "from": from_module,
                            "import": item,
                        })
                        
            # Extract classes
            classes = []
            for match in re.finditer(r"^class\s+(\w+)(?:\(([^)]*)\))?:", content, re.MULTILINE):
                class_name = match.group(1)
                parent_classes = []
                
                if match.group(2):
                    parent_classes = [c.strip() for c in match.group(2).split(",")]
                    
                classes.append({
                    "name": class_name,
                    "parents": parent_classes,
                    "position": match.start(),
                })
                
            # Extract functions
            functions = []
            for match in re.finditer(r"^def\s+(\w+)\s*\((.*?)\)(?:\s*->\s*([^:]+))?:", content, re.MULTILINE):
                func_name = match.group(1)
                params = match.group(2).strip()
                return_type = match.group(3).strip() if match.group(3) else None
                
                functions.append({
                    "name": func_name,
                    "params": params,
                    "return_type": return_type,
                    "position": match.start(),
                })
                
            return {
                "type": "python",
                "content": content,
                "docstring": docstring,
                "imports": imports,
                "classes": classes,
                "functions": functions,
                "metadata": {
                    "file_path": file_path,
                    "file_size": os.path.getsize(file_path),
                },
            }
        except Exception as e:
            logger.error(f"Error parsing Python file {file_path}: {str(e)}")
            return {
                "type": "python",
                "content": "",
                "error": str(e),
                "metadata": {
                    "file_path": file_path,
                },
            }
            
    def parse_javascript(self, file_path: str) -> Dict[str, Any]:
        """Parse JavaScript file.
        
        Args:
            file_path: Path to JavaScript file
            
        Returns:
            Dictionary with parsed content and structure
        """
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
                
            # Extract imports/requires
            imports = []
            # ES6 imports
            for match in re.finditer(r"import\s+(?:{([^}]+)}\s+from\s+)?['\"]([^'\"]+)['\"]", content):
                named_imports = []
                if match.group(1):
                    named_imports = [name.strip() for name in match.group(1).split(",")]
                    
                module = match.group(2)
                imports.append({
                    "type": "es6",
                    "module": module,
                    "named_imports": named_imports,
                })
                
            # CommonJS requires
            for match in re.finditer(r"(?:const|let|var)\s+(\w+)\s*=\s*require\(['\"]([^'\"]+)['\"]\)", content):
                variable = match.group(1)
                module = match.group(2)
                imports.append({
                    "type": "commonjs",
                    "module": module,
                    "variable": variable,
                })
                
            # Extract functions
            functions = []
            # Function declarations
            for match in re.finditer(r"function\s+(\w+)\s*\((.*?)\)", content):
                func_name = match.group(1)
                params = match.group(2).strip()
                
                functions.append({
                    "type": "declaration",
                    "name": func_name,
                    "params": params,
                    "position": match.start(),
                })
                
            # Arrow functions with assignments
            for match in re.finditer(r"(?:const|let|var)\s+(\w+)\s*=\s*(?:\((.*?)\)|(\w+))\s*=>\s*", content):
                func_name = match.group(1)
                params = match.group(2) or match.group(3) or ""
                
                functions.append({
                    "type": "arrow",
                    "name": func_name,
                    "params": params,
                    "position": match.start(),
                })
                
            # Extract classes
            classes = []
            for match in re.finditer(r"class\s+(\w+)(?:\s+extends\s+(\w+))?", content):
                class_name = match.group(1)
                parent_class = match.group(2)
                
                classes.append({
                    "name": class_name,
                    "extends": parent_class,
                    "position": match.start(),
                })
                
            return {
                "type": "javascript",
                "content": content,
                "imports": imports,
                "functions": functions,
                "classes": classes,
                "metadata": {
                    "file_path": file_path,
                    "file_size": os.path.getsize(file_path),
                },
            }
        except Exception as e:
            logger.error(f"Error parsing JavaScript file {file_path}: {str(e)}")
            return {
                "type": "javascript",
                "content": "",
                "error": str(e),
                "metadata": {
                    "file_path": file_path,
                },
            }
            
    def parse_docx(self, file_path: str) -> Dict[str, Any]:
        """Parse Word document.
        
        Args:
            file_path: Path to Word document
            
        Returns:
            Dictionary with parsed content
        """
        try:
            doc = docx.Document(file_path)
            
            # Extract text
            paragraphs = [p.text for p in doc.paragraphs]
            
            # Extract headings
            headings = []
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.style.name.startswith("Heading"):
                    level = int(paragraph.style.name.replace("Heading", ""))
                    headings.append({
                        "level": level,
                        "text": paragraph.text,
                        "position": i,
                    })
                    
            # Extract tables
            tables = []
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                    
                tables.append({
                    "index": i,
                    "data": table_data,
                })
                
            return {
                "type": "docx",
                "paragraphs": paragraphs,
                "headings": headings,
                "tables": tables,
                "metadata": {
                    "file_path": file_path,
                    "file_size": os.path.getsize(file_path),
                },
            }
        except Exception as e:
            logger.error(f"Error parsing Word document {file_path}: {str(e)}")
            return {
                "type": "docx",
                "content": "",
                "error": str(e),
                "metadata": {
                    "file_path": file_path,
                },
            }
            
    def parse_json(self, file_path: str) -> Dict[str, Any]:
        """Parse JSON file.
        
        Args:
            file_path: Path to JSON file
            
        Returns:
            Dictionary with parsed content
        """
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
                
            return {
                "type": "json",
                "data": data,
                "metadata": {
                    "file_path": file_path,
                    "file_size": os.path.getsize(file_path),
                },
            }
        except Exception as e:
            logger.error(f"Error parsing JSON file {file_path}: {str(e)}")
            return {
                "type": "json",
                "data": {},
                "error": str(e),
                "metadata": {
                    "file_path": file_path,
                },
            }
            
    def parse_yaml(self, file_path: str) -> Dict[str, Any]:
        """Parse YAML file.
        
        Args:
            file_path: Path to YAML file
            
        Returns:
            Dictionary with parsed content
        """
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                data = yaml.safe_load(f)
                
            return {
                "type": "yaml",
                "data": data,
                "metadata": {
                    "file_path": file_path,
                    "file_size": os.path.getsize(file_path),
                },
            }
        except Exception as e:
            logger.error(f"Error parsing YAML file {file_path}: {str(e)}")
            return {
                "type": "yaml",
                "data": {},
                                "error": str(e),
                "metadata": {
                    "file_path": file_path,
                },
            }


class KnowledgeExtractor:
    """Extracts structured knowledge from parsed documents."""

    def __init__(self):
        """Initialize knowledge extractor."""
        self.extractors = {
            "markdown": self.extract_from_markdown,
            "python": self.extract_from_python,
            "javascript": self.extract_from_javascript,
            "docx": self.extract_from_docx,
            "json": self.extract_from_json,
            "yaml": self.extract_from_yaml,
            "text": self.extract_from_text,
        }
        
    def extract(self, parsed_document: Dict[str, Any]) -> Dict[str, Any]:
        """Extract knowledge from parsed document.
        
        Args:
            parsed_document: Document parsed by DocumentParser
            
        Returns:
            Dictionary with extracted knowledge
        """
        doc_type = parsed_document.get("type", "text")
        
        if doc_type not in self.extractors:
            logger.warning(f"No knowledge extractor for {doc_type} documents")
            return self.extract_from_text(parsed_document)
            
        return self.extractors[doc_type](parsed_document)
        
    def extract_from_text(self, parsed_document: Dict[str, Any]) -> Dict[str, Any]:
        """Extract knowledge from plain text.
        
        Args:
            parsed_document: Parsed text document
            
        Returns:
            Dictionary with extracted knowledge
        """
        content = parsed_document.get("content", "")
        
        # Extract key concepts (simple implementation)
        words = re.findall(r'\b\w+\b', content.lower())
        word_freq = {}
        
        for word in words:
            if len(word) > 3:  # Skip short words
                word_freq[word] = word_freq.get(word, 0) + 1
                
        # Get top keywords
        keywords = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]
        
        return {
            "type": "text_knowledge",
            "keywords": [k[0] for k in keywords],
            "word_frequencies": word_freq,
            "source": parsed_document.get("metadata", {}).get("file_path", "unknown"),
        }
        
    def extract_from_markdown(self, parsed_document: Dict[str, Any]) -> Dict[str, Any]:
        """Extract knowledge from Markdown.
        
        Args:
            parsed_document: Parsed Markdown document
            
        Returns:
            Dictionary with extracted knowledge
        """
        # Extract concepts from headings
        concepts = {}
        headings = parsed_document.get("headings", [])
        
        for heading in headings:
            level = heading.get("level", 1)
            text = heading.get("text", "")
            
            if level <= 2:  # Main concepts are h1 and h2
                concepts[text] = {
                    "level": level,
                    "subconcepts": [],
                }
                
        # Add subconcepts
        current_concept = None
        
        for heading in headings:
            level = heading.get("level", 1)
            text = heading.get("text", "")
            
            if level <= 2:
                current_concept = text
            elif current_concept and level > 2:
                if current_concept in concepts:
                    concepts[current_concept]["subconcepts"].append({
                        "level": level,
                        "text": text,
                    })
                    
        # Extract code examples
        code_examples = []
        for code_block in parsed_document.get("code_blocks", []):
            code_examples.append({
                "language": code_block.get("language", "text"),
                "code": code_block.get("code", ""),
            })
            
        # Extract metadata
        metadata = parsed_document.get("front_matter", {})
        
        return {
            "type": "markdown_knowledge",
            "concepts": concepts,
            "code_examples": code_examples,
            "metadata": metadata,
            "source": parsed_document.get("metadata", {}).get("file_path", "unknown"),
        }
        
    def extract_from_python(self, parsed_document: Dict[str, Any]) -> Dict[str, Any]:
        """Extract knowledge from Python code.
        
        Args:
            parsed_document: Parsed Python document
            
        Returns:
            Dictionary with extracted knowledge
        """
        # Extract module purpose from docstring
        docstring = parsed_document.get("docstring", "")
        
        # Extract dependencies
        dependencies = []
        for imp in parsed_document.get("imports", []):
            from_module = imp.get("from")
            import_name = imp.get("import")
            
            if from_module:
                dependencies.append(f"{from_module}.{import_name}")
            else:
                dependencies.append(import_name)
                
        # Extract classes with methods
        classes = {}
        for cls in parsed_document.get("classes", []):
            class_name = cls.get("name", "")
            classes[class_name] = {
                "parents": cls.get("parents", []),
                "methods": [],
            }
            
        # Extract functions
        functions = {}
        for func in parsed_document.get("functions", []):
            func_name = func.get("name", "")
            functions[func_name] = {
                "params": func.get("params", ""),
                "return_type": func.get("return_type"),
            }
            
        return {
            "type": "python_knowledge",
            "module_purpose": docstring,
            "dependencies": dependencies,
            "classes": classes,
            "functions": functions,
            "source": parsed_document.get("metadata", {}).get("file_path", "unknown"),
        }
        
    def extract_from_javascript(self, parsed_document: Dict[str, Any]) -> Dict[str, Any]:
        """Extract knowledge from JavaScript code.
        
        Args:
            parsed_document: Parsed JavaScript document
            
        Returns:
            Dictionary with extracted knowledge
        """
        # Extract dependencies
        dependencies = []
        for imp in parsed_document.get("imports", []):
            if imp.get("type") == "es6":
                module = imp.get("module", "")
                named_imports = imp.get("named_imports", [])
                
                if named_imports:
                    for named_import in named_imports:
                        dependencies.append(f"{module}:{named_import}")
                else:
                    dependencies.append(module)
            else:  # CommonJS
                module = imp.get("module", "")
                variable = imp.get("variable", "")
                dependencies.append(f"{module} as {variable}")
                
        # Extract functions
        functions = {}
        for func in parsed_document.get("functions", []):
            func_name = func.get("name", "")
            functions[func_name] = {
                "type": func.get("type", "declaration"),
                "params": func.get("params", ""),
            }
            
        # Extract classes
        classes = {}
        for cls in parsed_document.get("classes", []):
            class_name = cls.get("name", "")
            classes[class_name] = {
                "extends": cls.get("extends"),
            }
            
        return {
            "type": "javascript_knowledge",
            "dependencies": dependencies,
            "functions": functions,
            "classes": classes,
            "source": parsed_document.get("metadata", {}).get("file_path", "unknown"),
        }
        
    def extract_from_docx(self, parsed_document: Dict[str, Any]) -> Dict[str, Any]:
        """Extract knowledge from Word document.
        
        Args:
            parsed_document: Parsed Word document
            
        Returns:
            Dictionary with extracted knowledge
        """
        # Extract concepts from headings (similar to Markdown)
        concepts = {}
        headings = parsed_document.get("headings", [])
        
        for heading in headings:
            level = heading.get("level", 1)
            text = heading.get("text", "")
            
            if level <= 2:  # Main concepts are h1 and h2
                concepts[text] = {
                    "level": level,
                    "subconcepts": [],
                }
                
        # Add subconcepts
        current_concept = None
        
        for heading in headings:
            level = heading.get("level", 1)
            text = heading.get("text", "")
            
            if level <= 2:
                current_concept = text
            elif current_concept and level > 2:
                if current_concept in concepts:
                    concepts[current_concept]["subconcepts"].append({
                        "level": level,
                        "text": text,
                    })
                    
        # Extract tables as structured data
        tables = parsed_document.get("tables", [])
        
        return {
            "type": "docx_knowledge",
            "concepts": concepts,
            "tables": tables,
            "paragraphs": parsed_document.get("paragraphs", []),
            "source": parsed_document.get("metadata", {}).get("file_path", "unknown"),
        }
        
    def extract_from_json(self, parsed_document: Dict[str, Any]) -> Dict[str, Any]:
        """Extract knowledge from JSON.
        
        Args:
            parsed_document: Parsed JSON document
            
        Returns:
            Dictionary with extracted knowledge
        """
        data = parsed_document.get("data", {})
        
        # Analyze structure
        structure = self._analyze_json_structure(data)
        
        return {
            "type": "json_knowledge",
            "data": data,
            "structure": structure,
            "source": parsed_document.get("metadata", {}).get("file_path", "unknown"),
        }
        
    def _analyze_json_structure(self, data: Any, path: str = "$") -> Dict[str, Any]:
        """Recursively analyze JSON structure.
        
        Args:
            data: JSON data
            path: Current JSON path
            
        Returns:
            Dictionary with structure information
        """
        if isinstance(data, dict):
            return {
                "type": "object",
                "properties": {
                    key: self._analyze_json_structure(value, f"{path}.{key}")
                    for key, value in data.items()
                },
                "path": path,
            }
        elif isinstance(data, list):
            if not data:
                return {
                    "type": "array",
                    "items": None,
                    "path": path,
                }
            
            # Analyze first item as representative
            return {
                "type": "array",
                "items": self._analyze_json_structure(data[0], f"{path}[0]"),
                "length": len(data),
                "path": path,
            }
        elif isinstance(data, str):
            return {
                "type": "string",
                "path": path,
            }
        elif isinstance(data, bool):
            return {
                "type": "boolean",
                "path": path,
            }
        elif isinstance(data, (int, float)):
            return {
                "type": "number",
                "path": path,
            }
        elif data is None:
            return {
                "type": "null",
                "path": path,
            }
        else:
            return {
                "type": str(type(data).__name__),
                "path": path,
            }
            
    def extract_from_yaml(self, parsed_document: Dict[str, Any]) -> Dict[str, Any]:
        """Extract knowledge from YAML.
        
        Args:
            parsed_document: Parsed YAML document
            
        Returns:
            Dictionary with extracted knowledge
        """
        data = parsed_document.get("data", {})
        
        # Analyze structure (reuse JSON structure analysis)
        structure = self._analyze_json_structure(data)
        
        return {
            "type": "yaml_knowledge",
            "data": data,
            "structure": structure,
            "source": parsed_document.get("metadata", {}).get("file_path", "unknown"),
        }


class TeachAndTrainInterface:
    """Interface for teaching Grace from various document types."""

    def __init__(self, config: Optional[Config] = None):
        """Initialize teach and train interface.
        
        Args:
            config: Configuration object. If None, uses default config.
        """
        self.config = config or Config()
        self.parser = DocumentParser()
        self.extractor = KnowledgeExtractor()
        self.knowledge_base = {}
        
    def teach(self, file_path: str) -> Dict[str, Any]:
        """Teach Grace from a single file.
        
        Args:
            file_path: Path to file
            
        Returns:
            Dictionary with extracted knowledge
        """
        logger.info(f"Teaching from file: {file_path}")
        
        # Parse document
        parsed_document = self.parser.parse_file(file_path)
        
        if "error" in parsed_document:
            logger.error(f"Error parsing {file_path}: {parsed_document['error']}")
            return {
                "success": False,
                "error": parsed_document["error"],
                "file_path": file_path,
            }
            
        # Extract knowledge
        knowledge = self.extractor.extract(parsed_document)
        
        # Store in knowledge base
        self.knowledge_base[file_path] = knowledge
        
        return {
            "success": True,
            "knowledge": knowledge,
            "file_path": file_path,
        }
        
    def train_on_directory(self, directory_path: str, recursive: bool = True) -> Dict[str, Any]:
        """Train Grace on all files in a directory.
        
        Args:
            directory_path: Path to directory
            recursive: Whether to recursively process subdirectories
            
        Returns:
            Dictionary with training results
        """
        logger.info(f"Training on directory: {directory_path}")
        
        results = {
            "success": True,
            "processed_files": 0,
            "failed_files": 0,
            "knowledge_items": 0,
            "errors": [],
        }
        
        if not os.path.exists(directory_path):
            results["success"] = False
            results["errors"].append(f"Directory not found: {directory_path}")
            return results
            
        for root, dirs, files in os.walk(directory_path):
            for file in files:
                file_path = os.path.join(root, file)
                
                try:
                    teach_result = self.teach(file_path)
                    
                    if teach_result["success"]:
                        results["processed_files"] += 1
                        results["knowledge_items"] += 1
                    else:
                        results["failed_files"] += 1
                        results["errors"].append({
                            "file": file_path,
                            "error": teach_result.get("error", "Unknown error"),
                        })
                except Exception as e:
                    results["failed_files"] += 1
                    results["errors"].append({
                        "file": file_path,
                        "error": str(e),
                    })
                    
            if not recursive:
                break
                
        return results
        
    def save_knowledge_base(self, output_path: str) -> bool:
        """Save knowledge base to file.
        
        Args:
            output_path: Path to save knowledge base
            
        Returns:
            True if successful, False otherwise
        """
        try:
                        with open(output_path, "w", encoding="utf-8") as f:
                json.dump(self.knowledge_base, f, indent=2)
                
            logger.info(f"Knowledge base saved to {output_path}")
            return True
        except Exception as e:
            logger.error(f"Error saving knowledge base: {str(e)}")
            return False
            
    def load_knowledge_base(self, input_path: str) -> bool:
        """Load knowledge base from file.
        
        Args:
            input_path: Path to load knowledge base from
            
        Returns:
            True if successful, False otherwise
        """
        try:
            with open(input_path, "r", encoding="utf-8") as f:
                self.knowledge_base = json.load(f)
                
            logger.info(f"Knowledge base loaded from {input_path}")
            return True
        except Exception as e:
            logger.error(f"Error loading knowledge base: {str(e)}")
            return False
            
    def query_knowledge(self, query: str) -> List[Dict[str, Any]]:
        """Query knowledge base.
        
        Args:
            query: Query string
            
        Returns:
            List of matching knowledge items
        """
        results = []
        
        # Simple keyword matching for now
        query_terms = set(query.lower().split())
        
        for file_path, knowledge in self.knowledge_base.items():
            # Convert knowledge to string for simple matching
            knowledge_str = json.dumps(knowledge).lower()
            
            # Check if all query terms are in the knowledge
            if all(term in knowledge_str for term in query_terms):
                results.append({
                    "file_path": file_path,
                    "knowledge": knowledge,
                    "relevance": 1.0,  # Simple relevance score
                })
                
        # Sort by relevance
        results.sort(key=lambda x: x["relevance"], reverse=True)
        
        return results
        
    def get_knowledge_summary(self) -> Dict[str, Any]:
        """Get summary of knowledge base.
        
        Returns:
            Dictionary with knowledge base summary
        """
        summary = {
            "total_files": len(self.knowledge_base),
            "knowledge_types": {},
            "sources": [],
        }
        
        for file_path, knowledge in self.knowledge_base.items():
            knowledge_type = knowledge.get("type", "unknown")
            
            if knowledge_type not in summary["knowledge_types"]:
                summary["knowledge_types"][knowledge_type] = 0
                
            summary["knowledge_types"][knowledge_type] += 1
            summary["sources"].append(file_path)
            
        return summary

